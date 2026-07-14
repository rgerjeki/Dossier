"""UI tests for the main window. Skipped entirely when PySide6 is not installed,
so the core engine suite stays green in a minimal environment.

These drive the window with a fake collector: no Qt event loop threading, no
network. The background worker is exercised separately; here we verify the window
maps a CollectorRun into the table and that curation edits flow back into the
underlying Finding objects.
"""

from __future__ import annotations

import os

import pytest

pytest.importorskip("PySide6")

os.environ.setdefault("QT_QPA_PLATFORM", "offscreen")

from PySide6.QtCore import Qt, Signal  # noqa: E402
from PySide6.QtWidgets import QApplication, QWidget  # noqa: E402

from dossier.collectors.base import Collector, CollectorRun  # noqa: E402
from dossier.models import Confidence, Finding, FindingStatus, FindingType  # noqa: E402
from dossier.ui.window import (  # noqa: E402
    COL_ANALYST,
    COL_INCLUDE,
    COL_NOTES,
    COL_VALUE,
    MainWindow,
)


@pytest.fixture(scope="module")
def qapp():
    app = QApplication.instance() or QApplication([])
    yield app


class _FakeCollector(Collector):
    name = "Fake"

    def __init__(self, run: CollectorRun) -> None:
        self._run = run

    def collect(self, target: str) -> CollectorRun:
        return self._run


class _FakeBridge:
    def __init__(self, editor: _FakeReportEditor) -> None:
        self._editor = editor

    def onInput(self, html: str) -> None:  # mirrors the real editor's JS bridge
        self._editor._content = html
        self._editor.contentEdited.emit()


class _FakeReportEditor(QWidget):
    """A no-Chromium stand-in for the web report editor, for fast/stable tests."""

    contentEdited = Signal()

    def __init__(self) -> None:
        super().__init__()
        self._content = ""
        self._dynamic: dict[str, str] = {}
        self._bridge = _FakeBridge(self)

    def set_content(self, html: str) -> None:
        self._content = html
        self._dynamic = {}

    def update_dynamic(self, parts: dict) -> None:
        self._dynamic = dict(parts)
        self.contentEdited.emit()

    def content(self) -> str:
        # Approximates the real editor: base document plus the live findings parts.
        return self._content + "".join(self._dynamic.values())

    def scroll_to_anchor(self, name: str) -> None:
        pass

    def print_pdf(self, path, on_done=None) -> None:
        from pathlib import Path

        Path(path).write_bytes(b"%PDF-1.4 fake")
        if on_done is not None:
            on_done(True)


def _run_with(*findings: Finding, ok: bool = True, message: str = "done") -> CollectorRun:
    return CollectorRun("Fake", findings=list(findings), ok=ok, message=message)


def _acct(value: str, source: str = "Maigret (GitHub)", **kw) -> Finding:
    return Finding(type=FindingType.ACCOUNT, value=value, source=source, **kw)


def _make_window(collector=None) -> MainWindow:
    # Always inject the fake report editor so tests never start QtWebEngine.
    return MainWindow(
        collector=collector, report_editor_factory=_FakeReportEditor
    )


def _window_with_findings(qapp, *findings: Finding) -> MainWindow:
    window = _make_window(collector=_FakeCollector(_run_with(*findings)))
    window._input.setText("subject")
    window._on_done(_run_with(*findings))
    return window


def test_window_populates_table_from_run(qapp) -> None:
    window = _window_with_findings(
        qapp,
        _acct("https://github.com/ex"),
        _acct("Twitter (ex)", "Maigret (Twitter)", status=FindingStatus.UNREACHABLE),
    )
    assert window._table.rowCount() == 2
    assert window._table.item(0, COL_VALUE).text() == "https://github.com/ex"
    assert "2 findings" in window._info.text()
    assert "0 included" in window._info.text()


def test_toggle_include_updates_finding(qapp) -> None:
    finding = _acct("https://github.com/ex")
    window = _window_with_findings(qapp, finding)

    assert finding.included is False
    window._table.item(0, COL_INCLUDE).setCheckState(Qt.CheckState.Checked)
    assert finding.included is True
    assert "1 included" in window._info.text()


def test_edit_notes_updates_finding(qapp) -> None:
    finding = Finding(type=FindingType.ACCOUNT, value="x", source="s")
    window = _window_with_findings(qapp, finding)

    window._table.item(0, COL_NOTES).setText("looks legit, matches other handles")
    assert finding.notes == "looks legit, matches other handles"


def test_analyst_confidence_dropdown_updates_finding(qapp) -> None:
    finding = Finding(type=FindingType.ACCOUNT, value="x", source="s")
    window = _window_with_findings(qapp, finding)

    combo = window._table.cellWidget(0, COL_ANALYST)  # always-visible dropdown
    combo.setCurrentText(Confidence.HIGH.value)
    assert finding.analyst_confidence is Confidence.HIGH


def test_filter_hides_non_matching_rows(qapp) -> None:
    window = _window_with_findings(
        qapp,
        _acct("https://github.com/ex"),
        _acct("https://reddit.com/u/ex", "Maigret (Reddit)"),
    )
    window._filter.setText("github")
    assert not window._table.isRowHidden(0)
    assert window._table.isRowHidden(1)
    window._filter.setText("")
    assert not window._table.isRowHidden(1)


def test_header_click_sorts_findings(qapp) -> None:
    window = _window_with_findings(qapp, _acct("ccc"), _acct("aaa"), _acct("bbb"))
    window._on_header_clicked(COL_VALUE)  # ascending by Value
    assert [window._table.item(r, COL_VALUE).text() for r in range(3)] == [
        "aaa",
        "bbb",
        "ccc",
    ]
    window._on_header_clicked(COL_VALUE)  # toggle descending
    assert window._table.item(0, COL_VALUE).text() == "ccc"




def test_save_and_open_case_round_trip(qapp, tmp_path) -> None:
    finding = _acct("https://github.com/ex")
    window = _window_with_findings(qapp, finding)
    # curate before saving
    window._table.item(0, COL_INCLUDE).setCheckState(Qt.CheckState.Checked)
    window._table.item(0, COL_NOTES).setText("kept")

    path = tmp_path / "case.json"
    window.save_case(str(path))
    assert path.exists()

    reopened = _make_window(collector=_FakeCollector(_run_with()))
    reopened.open_case(str(path))
    assert reopened._table.rowCount() == 1
    assert reopened._case.included_findings()[0].notes == "kept"
    assert reopened._input.text() == "subject"


def test_create_case_adopts_and_saves(qapp, tmp_path) -> None:
    from dossier.case import Case, SubjectType

    window = _make_window(collector=_FakeCollector(_run_with()))
    case = Case(subject="ada", subject_type=SubjectType.USERNAME, title="Ada case")
    path = tmp_path / "ada.json"
    window.create_case(case, str(path))

    assert window._case is case
    assert window._case_path == str(path)
    assert window._dirty is False
    assert path.exists()
    assert window._input.text() == "ada"


def test_create_case_without_path_is_unsaved(qapp) -> None:
    from dossier.case import Case, SubjectType

    window = _make_window(collector=_FakeCollector(_run_with()))
    window.create_case(Case(subject="x", subject_type=SubjectType.USERNAME, title="X"))
    assert window._case is not None
    assert window._case_path is None
    assert window._dirty is True


def test_new_case_dialog_builds_case(qapp) -> None:
    from dossier.case import SubjectType
    from dossier.ui.window import NewCaseDialog

    dialog = NewCaseDialog()
    dialog._name.setText("My Investigation")
    dialog._subject.setText("target_user")
    case = dialog.build_case()
    assert case.title == "My Investigation"
    assert case.subject == "target_user"
    assert case.subject_type is SubjectType.USERNAME


def test_failed_run_shows_note(qapp) -> None:
    failed = _run_with(ok=False, message="Maigret is not installed.")
    window = _make_window(collector=_FakeCollector(failed))
    window._input.setText("subject")
    window._on_done(failed)
    assert window._table.rowCount() == 0
    assert window._status.text().startswith("Note: ")


def test_empty_username_does_not_run(qapp) -> None:
    window = _make_window(collector=_FakeCollector(_run_with()))
    window._input.setText("   ")
    window.run_collection()
    assert window._table.rowCount() == 0
    assert "Enter a value" in window._status.text()


def test_report_view_populates_outline_and_preview(qapp) -> None:
    window = _window_with_findings(qapp, _acct("https://github.com/ex", included=True))
    window.show_report_view()
    assert window._stack.currentIndex() == 1
    titles = [window._outline.item(i).text() for i in range(window._outline.count())]
    assert "Overview" in titles
    assert "Sources" in titles
    assert "github.com/ex" in window._report_editor().content()


def test_report_reflects_curation_live(qapp) -> None:
    a = _acct("https://github.com/alpha", included=True)
    b = _acct("https://reddit.com/u/beta", "Maigret (Reddit)", included=False)
    window = _window_with_findings(qapp, a, b)

    window.show_report_view()
    assert "github.com/alpha" in window._report_editor().content()
    assert "reddit.com/u/beta" not in window._report_editor().content()

    # Include the second finding via the table; the report updates live.
    window._table.item(1, COL_INCLUDE).setCheckState(Qt.CheckState.Checked)
    assert "reddit.com/u/beta" in window._report_editor().content()


def test_hand_edits_are_not_overwritten_by_curation(qapp) -> None:
    finding = _acct("https://github.com/ex", included=True)
    window = _window_with_findings(qapp, finding)
    window.show_report_view()
    window._report_editor()._bridge.onInput("<p>Hand-written analysis.</p>")

    # A curation change updates the findings but must not wipe the analyst text.
    window._table.item(0, COL_INCLUDE).setCheckState(Qt.CheckState.Unchecked)
    assert "Hand-written analysis." in window._report_editor().content()


def test_editing_report_updates_case(qapp) -> None:
    window = _window_with_findings(qapp, _acct("https://github.com/ex", included=True))
    window.show_report_view()
    # Simulate the analyst typing into the document (the editor mirrors edits
    # back over the bridge, which the window persists onto the case).
    window._report_editor()._bridge.onInput("<p>Analyst note typed into the doc.</p>")
    assert "Analyst note typed into the doc." in window._case.report_html


def test_saved_edited_report_reopens(qapp, tmp_path) -> None:
    window = _window_with_findings(qapp, _acct("https://github.com/ex", included=True))
    window.show_report_view()
    window._report_editor()._bridge.onInput("<p>A unique hand-typed sentence.</p>")
    path = tmp_path / "case.json"
    window.save_case(str(path))

    reopened = _make_window(collector=_FakeCollector(_run_with()))
    reopened.show_report_view()
    reopened.open_case(str(path))
    assert "A unique hand-typed sentence." in reopened._report_editor().content()


def test_include_all_and_exclude_all(qapp) -> None:
    window = _window_with_findings(
        qapp, _acct("a"), _acct("b"), _acct("c")
    )
    window._set_all_included(True)
    assert len(window._case.included_findings()) == 3
    window._set_all_included(False)
    assert len(window._case.included_findings()) == 0


def test_open_refreshes_report_when_on_report_view(qapp, tmp_path) -> None:
    # Build and save a case with a distinctive finding.
    src = _window_with_findings(qapp, _acct("https://github.com/unique_handle", included=True))
    path = tmp_path / "c.json"
    src.save_case(str(path))

    # A second window sitting on the report view opens the file: preview must update.
    window = _make_window(collector=_FakeCollector(_run_with()))
    window.show_report_view()
    window.open_case(str(path))
    assert "unique_handle" in window._report_editor().content()


def test_subject_type_preselects_collector(qapp, tmp_path) -> None:
    from dossier.case import Case, SubjectType
    from dossier.ui.window import CHOICE_EMAIL

    window = _make_window(collector=_FakeCollector(_run_with()))
    window.create_case(
        Case(subject="a@b.com", subject_type=SubjectType.EMAIL, title="E"),
        str(tmp_path / "e.json"),
    )
    assert window._choice.currentText() == CHOICE_EMAIL


def test_view_switching(qapp) -> None:
    window = _window_with_findings(qapp, _acct("a", included=True))
    window.show_report_view()
    assert window._stack.currentIndex() == 1
    window.show_collection_view()
    assert window._stack.currentIndex() == 0


def test_prewarm_creates_editor(qapp) -> None:
    window = _make_window(collector=_FakeCollector(_run_with()))
    assert window._editor is None
    window.prewarm_report_editor()
    assert window._editor is not None  # built ahead of first Report-tab click


def test_collector_choice_selects_collector(qapp) -> None:
    from dossier.collectors.courtlistener import CourtListenerCollector
    from dossier.collectors.email import EmailCollector
    from dossier.collectors.github import GitHubCollector
    from dossier.collectors.keybase import KeybaseCollector
    from dossier.collectors.searchkit import SearchKitCollector
    from dossier.collectors.sec import SECCollector
    from dossier.collectors.usernames import MaigretCollector
    from dossier.ui.window import (
        CHOICE_COURT,
        CHOICE_EMAIL,
        CHOICE_GITHUB,
        CHOICE_KEYBASE,
        CHOICE_SEARCHKIT,
        CHOICE_SEC,
        CHOICE_USERNAME,
    )

    window = _make_window()  # no injected collector -> honors the choice
    for choice, cls in (
        (CHOICE_USERNAME, MaigretCollector),
        (CHOICE_GITHUB, GitHubCollector),
        (CHOICE_KEYBASE, KeybaseCollector),
        (CHOICE_EMAIL, EmailCollector),
        (CHOICE_SEC, SECCollector),
        (CHOICE_COURT, CourtListenerCollector),
        (CHOICE_SEARCHKIT, SearchKitCollector),
    ):
        window._choice.setCurrentText(choice)
        assert isinstance(window._get_collector(), cls)
        assert window._input.placeholderText()


def test_remove_findings_deletes_from_case(qapp) -> None:
    keep = _acct("https://github.com/keep", included=True)
    drop = _acct("https://reddit.com/u/drop", "Maigret (Reddit)", included=True)
    window = _window_with_findings(qapp, keep, drop)

    window._remove_findings([drop])  # the no-dialog core
    assert [f.value for f in window._case.findings] == ["https://github.com/keep"]
    assert window._table.rowCount() == 1
    assert window._dirty is True

    # The report's dynamic parts (what patches the live report) no longer mention it.
    from dossier.report.render import render_dynamic_parts

    parts = " ".join(render_dynamic_parts(window._case).values())
    assert "reddit.com/u/drop" not in parts


def test_remove_excluded_prunes_noise(qapp) -> None:
    window = _window_with_findings(
        qapp,
        _acct("kept", included=True),
        _acct("noise1", included=False),
        _acct("noise2", included=False),
    )
    excluded = [f for f in window._case.findings if not f.included]
    window._remove_findings(excluded)  # what "Remove excluded" does after confirm
    assert len(window._case.findings) == 1
    assert window._case.findings[0].value == "kept"


def test_github_finding_is_enrichable(qapp) -> None:
    from dossier.ui.window import _github_username

    # A Maigret-style GitHub hit can be enriched to a username.
    maigret_hit = Finding(
        type=FindingType.ACCOUNT,
        value="https://github.com/rgerjeki",
        source="Maigret (GitHub)",
        source_url="https://github.com/rgerjeki",
    )
    assert _github_username(maigret_hit) == "rgerjeki"

    # The GitHub collector's own output is not re-enriched.
    own = Finding(
        type=FindingType.ACCOUNT,
        value="GitHub profile: https://github.com/rgerjeki",
        source="GitHub",
        source_url="https://github.com/rgerjeki",
    )
    assert _github_username(own) is None

    # A non-profile GitHub path (e.g. search) is not treated as a username.
    other = Finding(type=FindingType.LINK, value="x", source="Search-kit",
                    source_url="https://github.com/search?q=x")
    assert _github_username(other) is None


def test_collector_dropdown_is_grouped(qapp) -> None:
    window = _make_window()
    model = window._choice.model()
    headers = []
    choices = []
    for row in range(model.rowCount()):
        item = model.item(row)
        if item.flags() & Qt.ItemFlag.ItemIsSelectable:
            choices.append(item.text())
        else:
            headers.append(item.text())  # non-selectable section header
    assert "Specialized" in headers  # niche collectors sit apart
    assert "Keybase (username)" in choices
    assert len(choices) == 7  # all collectors still reachable


def test_export_docx_writes_file(qapp, tmp_path) -> None:
    pytest.importorskip("htmldocx")
    finding = _acct("https://github.com/ex", included=True)
    window = _window_with_findings(qapp, finding)
    out = tmp_path / "report.docx"
    window.export_docx(str(out))
    assert out.exists() and out.stat().st_size > 0

    from docx import Document

    doc = Document(out)
    text = "\n".join(p.text for p in doc.paragraphs)
    for table in doc.tables:  # findings render into tables
        for row in table.rows:
            text += "\n" + "\t".join(cell.text for cell in row.cells)
    assert "https://github.com/ex" in text


def test_export_without_case_is_a_noop(qapp, tmp_path) -> None:
    window = _make_window(collector=_FakeCollector(_run_with()))
    out = tmp_path / "nope.pdf"
    window.export_pdf(str(out))
    assert not out.exists()
    assert "No case" in window._status.text()


def test_dirty_set_on_change_and_cleared_on_save(qapp, tmp_path) -> None:
    window = _window_with_findings(qapp, _acct("a", included=True))
    assert window._dirty is True  # collecting created unsaved state
    window.save_case(str(tmp_path / "c.json"))
    assert window._dirty is False
    assert window._case_path == str(tmp_path / "c.json")
    # a further edit marks dirty again
    window._table.item(0, COL_INCLUDE).setCheckState(Qt.CheckState.Unchecked)
    assert window._dirty is True


def test_autosave_writes_to_current_path(qapp, tmp_path) -> None:
    window = _window_with_findings(qapp, _acct("a", included=True))
    path = tmp_path / "auto.json"
    window.save_case(str(path))
    window._table.item(0, COL_INCLUDE).setCheckState(Qt.CheckState.Unchecked)
    assert window._dirty is True
    window._autosave()
    assert window._dirty is False
    from dossier.case import Case

    reloaded = Case.load(str(path))
    assert reloaded.included_findings() == []  # the unchecked edit was persisted


def test_save_returns_false_without_case(qapp) -> None:
    window = _make_window(collector=_FakeCollector(_run_with()))
    assert window.save_case("ignored.json") is False


def test_save_writes_to_current_file_without_dialog(qapp, tmp_path) -> None:
    window = _window_with_findings(qapp, _acct("a", included=True))
    path = tmp_path / "c.json"
    window.save_case(str(path))  # first save establishes the current file
    assert window._case_path == str(path)

    window._table.item(0, COL_INCLUDE).setCheckState(Qt.CheckState.Unchecked)
    assert window._dirty is True
    assert window.save_case() is True  # no path passed -> uses current file
    assert window._dirty is False

    from dossier.case import Case

    assert Case.load(str(path)).included_findings() == []
